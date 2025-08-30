import streamlit as st
import io
from datetime import datetime
import os

# Try importing required libraries with proper error handling
try:
    import PyPDF2
    PDF_LIBRARY_AVAILABLE = True
    PDF_LIBRARY = "PyPDF2"
except ImportError:
    try:
        import pdfplumber
        PDF_LIBRARY_AVAILABLE = True
        PDF_LIBRARY = "pdfplumber"
    except ImportError:
        PDF_LIBRARY_AVAILABLE = False
        PDF_LIBRARY = None

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_LIBRARY_AVAILABLE = True
except ImportError:
    DOCX_LIBRARY_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# --- Page Config ---
st.set_page_config(
    page_title="📄↔️📝 PDF ↔️ Word Converter",
    page_icon="🔄",
    layout="wide"
)

st.title("📄↔️📝 PDF ↔️ Word Converter")
st.markdown("Convert between PDF and Word documents with ease!")

# Check for required libraries
missing_libraries = []
if not PDF_LIBRARY_AVAILABLE:
    missing_libraries.extend(["PyPDF2", "pdfplumber"])
if not DOCX_LIBRARY_AVAILABLE:
    missing_libraries.append("python-docx")
if not REPORTLAB_AVAILABLE:
    missing_libraries.append("reportlab")

if missing_libraries:
    st.error(f"""
    **Missing Required Libraries!**
    
    Please add the following to your `requirements.txt` file:
    
    ```
    streamlit
    PyPDF2
    pdfplumber
    python-docx
    reportlab
    ```
    
    **Missing libraries:** {', '.join(missing_libraries)}
    """)
    st.stop()

# --- Conversion Mode Selection ---
st.subheader("🔄 Select Conversion Mode")
conversion_mode = st.selectbox(
    "Choose conversion type:",
    ["📄➡️📝 PDF to Word", "📝➡️📄 Word to PDF"],
    help="Select whether you want to convert PDF to Word or Word to PDF"
)

# --- Helper Functions for PDF to Word ---
def extract_text_pypdf2(pdf_file):
    """Extract text from PDF using PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append({
                        'page_number': page_num + 1,
                        'text': page_text
                    })
            except Exception as e:
                st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                continue
        
        return text_content, len(pdf_reader.pages)
    except Exception as e:
        st.error(f"Error reading PDF with PyPDF2: {str(e)}")
        return None, 0

def extract_text_pdfplumber(pdf_file):
    """Extract text from PDF using pdfplumber"""
    try:
        import pdfplumber
        text_content = []
        
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content.append({
                            'page_number': page_num + 1,
                            'text': page_text
                        })
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            return text_content, len(pdf.pages)
    except Exception as e:
        st.error(f"Error reading PDF with pdfplumber: {str(e)}")
        return None, 0

def create_word_document(text_content, original_filename, formatting_options):
    """Create a Word document from extracted text"""
    try:
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Converted from: {original_filename}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add conversion info
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f'Converted on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        info_run.italic = True
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add a page break
        doc.add_page_break()
        
        # Process each page
        for page_data in text_content:
            page_num = page_data['page_number']
            page_text = page_data['text']
            
            # Add page header if requested
            if formatting_options['include_page_numbers']:
                page_header = doc.add_heading(f'Page {page_num}', level=2)
                page_header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Split text into paragraphs
            paragraphs = page_text.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Clean up the text
                    cleaned_text = para_text.strip().replace('\n', ' ')
                    
                    # Add paragraph
                    para = doc.add_paragraph(cleaned_text)
                    
                    # Apply formatting
                    if formatting_options['font_size'] != 11:
                        for run in para.runs:
                            run.font.size = Pt(formatting_options['font_size'])
                    
                    # Add spacing
                    para.space_after = Pt(6)
            
            # Add page break between PDF pages if requested
            if formatting_options['page_breaks'] and page_num < len(text_content):
                doc.add_page_break()
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

# --- Helper Functions for Word to PDF ---
def extract_text_from_docx(docx_file):
    """Extract text from Word document"""
    try:
        # Reset file pointer to beginning
        docx_file.seek(0)
        
        doc = Document(docx_file)
        text_content = []
        
        # Check if document has any content
        if not doc.paragraphs:
            st.warning("⚠️ The Word document appears to be empty or could not be read.")
            return None
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine if it's a heading
                is_heading = para.style.name.startswith('Heading')
                heading_level = 0
                if is_heading:
                    try:
                        heading_level = int(para.style.name.split()[-1])
                    except:
                        heading_level = 1
                
                text_content.append({
                    'text': para.text.strip(),
                    'is_heading': is_heading,
                    'heading_level': heading_level,
                    'style': para.style.name
                })
        
        if not text_content:
            st.warning("⚠️ No readable text found in the Word document.")
            return None
            
        return text_content
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        st.error("Make sure you're uploading a valid .docx file (not .doc or other formats)")
        return None

def create_pdf_from_text(text_content, original_filename, pdf_options):
    """Create PDF from text content using ReportLab"""
    try:
        buffer = io.BytesIO()
        
        # Set page size
        page_size = A4 if pdf_options['page_size'] == 'A4' else letter
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=page_size,
            rightMargin=pdf_options['margin'],
            leftMargin=pdf_options['margin'],
            topMargin=pdf_options['margin'],
            bottomMargin=pdf_options['margin']
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph(f"Converted from: {original_filename}", title_style))
        
        # Add conversion info
        info_style = ParagraphStyle(
            'CustomInfo',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=20,
            alignment=1,  # Center alignment
            textColor='grey'
        )
        story.append(Paragraph(f"Converted on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", info_style))
        story.append(Spacer(1, 20))
        
        # Process content
        for item in text_content:
            text = item['text']
            
            if item['is_heading']:
                # Use heading style
                heading_level = item['heading_level']
                if heading_level == 1:
                    style = styles['Heading1']
                elif heading_level == 2:
                    style = styles['Heading2']
                else:
                    style = styles['Heading3']
                
                # Customize heading style
                heading_style = ParagraphStyle(
                    f'CustomHeading{heading_level}',
                    parent=style,
                    fontSize=pdf_options['font_size'] + (4 - heading_level) * 2,
                    spaceAfter=12
                )
                story.append(Paragraph(text, heading_style))
            else:
                # Use normal paragraph style
                para_style = ParagraphStyle(
                    'CustomParagraph',
                    parent=styles['Normal'],
                    fontSize=pdf_options['font_size'],
                    spaceAfter=pdf_options['line_spacing'],
                    alignment=0  # Left alignment
                )
                story.append(Paragraph(text, para_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None

# --- Main App Interface Based on Selected Mode ---
st.divider()

if conversion_mode == "📄➡️📝 PDF to Word":
    # PDF to Word Conversion Interface
    st.header("📄➡️📝 PDF to Word Conversion")
    
    col1, col2 = st.columns([2, 1])
    
    with col2:
        st.info(f"📚 Using **{PDF_LIBRARY}** for PDF processing")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type=["pdf"],
        help="Upload a PDF file to convert to Word format",
        key="pdf_upload"
    )
    
    if uploaded_file is not None:
        # Display file info
        with st.expander("📋 File Information", expanded=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Filename", uploaded_file.name)
            with col2:
                st.metric("Size", f"{uploaded_file.size / 1024:.1f} KB")
            with col3:
                st.metric("Type", "PDF")
        
        # Formatting options
        st.subheader("🎨 Word Document Formatting Options")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            include_page_numbers = st.checkbox("Include page numbers", value=True, key="pdf_page_nums")
        
        with col2:
            page_breaks = st.checkbox("Page breaks between PDF pages", value=True, key="pdf_breaks")
        
        with col3:
            font_size = st.selectbox("Font size", [9, 10, 11, 12, 14, 16], index=2, key="pdf_font")
        
        with col4:
            extraction_method = st.selectbox(
                "Extraction method",
                ["Auto", "PyPDF2", "pdfplumber"],
                help="Choose PDF text extraction method",
                key="pdf_method"
            )
        
        formatting_options = {
            'include_page_numbers': include_page_numbers,
            'page_breaks': page_breaks,
            'font_size': font_size
        }
        
        # Convert button
        if st.button("🔄 Convert PDF to Word", type="primary", use_container_width=True, key="pdf_convert"):
            with st.spinner("Converting PDF to Word... This may take a moment."):
                # Reset file pointer
                uploaded_file.seek(0)
                
                # Extract text based on method selection
                text_content = None
                total_pages = 0
                
                if extraction_method == "PyPDF2" or (extraction_method == "Auto" and PDF_LIBRARY == "PyPDF2"):
                    text_content, total_pages = extract_text_pypdf2(uploaded_file)
                elif extraction_method == "pdfplumber" or (extraction_method == "Auto" and PDF_LIBRARY == "pdfplumber"):
                    uploaded_file.seek(0)
                    text_content, total_pages = extract_text_pdfplumber(uploaded_file)
                
                if text_content and len(text_content) > 0:
                    st.success(f"✅ Successfully extracted text from {len(text_content)} pages out of {total_pages} total pages")
                    
                    # Preview extracted text
                    with st.expander("👁️ Preview Extracted Text", expanded=False):
                        preview_text = ""
                        char_count = 0
                        
                        for page_data in text_content[:3]:
                            page_preview = page_data['text'][:500]
                            preview_text += f"**Page {page_data['page_number']}:**\n{page_preview}...\n\n"
                            char_count += len(page_data['text'])
                        
                        st.text_area("Text Preview", preview_text, height=200, disabled=True, key="pdf_preview")
                        st.info(f"Total characters extracted: {char_count:,}")
                    
                    # Create Word document
                    original_filename = uploaded_file.name.replace('.pdf', '')
                    word_doc = create_word_document(text_content, original_filename, formatting_options)
                    
                    if word_doc:
                        output_filename = f"{original_filename}_converted.docx"
                        
                        st.download_button(
                            label="📥 Download Word Document",
                            data=word_doc,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True,
                            key="pdf_download"
                        )
                        
                        st.success("🎉 PDF to Word conversion completed successfully!")
                else:
                    st.warning("⚠️ Could not extract readable text from the PDF.")

else:  # Word to PDF Conversion
    st.header("📝➡️📄 Word to PDF Conversion")
    
    col1, col2 = st.columns([2, 1])
    
    with col2:
        st.info("🔧 Using **ReportLab** for PDF generation")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a Word document",
        type=["docx", "doc"],
        help="Upload a Word document (.docx or .doc) to convert to PDF format",
        key="word_upload"
    )
    
    # Debug information
    if uploaded_file is not None:
        st.success(f"✅ File uploaded successfully: {uploaded_file.name}")
    else:
        st.info("📤 Please select a Word document (.docx or .doc file)")
    
    if uploaded_file is not None:
        # Display file info
        with st.expander("📋 File Information", expanded=True):
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Filename", uploaded_file.name)
            with col2:
                st.metric("Size", f"{uploaded_file.size / 1024:.1f} KB")
            with col3:
                st.metric("Type", "DOCX")
        
        # PDF formatting options
        st.subheader("🎨 PDF Formatting Options")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            page_size = st.selectbox("Page size", ["A4", "Letter"], index=0, key="word_page_size")
        
        with col2:
            font_size = st.selectbox("Font size", [9, 10, 11, 12, 14, 16], index=2, key="word_font")
        
        with col3:
            line_spacing = st.selectbox("Line spacing", [6, 8, 10, 12, 15], index=2, key="word_spacing")
        
        with col4:
            margin = st.selectbox("Margins (inches)", [0.5, 0.75, 1.0, 1.25], index=2, key="word_margin")
        
        pdf_options = {
            'page_size': page_size,
            'font_size': font_size,
            'line_spacing': line_spacing,
            'margin': margin * inch
        }
        
        # Convert button
        if st.button("🔄 Convert Word to PDF", type="primary", use_container_width=True, key="word_convert"):
            with st.spinner("Converting Word to PDF... This may take a moment."):
                try:
                    # Reset file pointer
                    uploaded_file.seek(0)
                    
                    # Debug: Show file details
                    st.info(f"Processing file: {uploaded_file.name} ({uploaded_file.size} bytes)")
                    
                    # Extract text from Word document
                    text_content = extract_text_from_docx(uploaded_file)
                    
                    if text_content and len(text_content) > 0:
                        st.success(f"✅ Successfully extracted content from Word document ({len(text_content)} paragraphs)")
                        
                        # Preview extracted text
                        with st.expander("👁️ Preview Extracted Text", expanded=False):
                            preview_content = []
                            char_count = 0
                            
                            for item in text_content[:10]:  # Show first 10 items
                                text_type = "HEADING" if item['is_heading'] else "PARAGRAPH"
                                preview_content.append(f"[{text_type}] {item['text'][:200]}...")
                                char_count += len(item['text'])
                            
                            preview_text = "\n\n".join(preview_content)
                            st.text_area("Content Preview", preview_text, height=300, disabled=True, key="word_preview")
                            st.info(f"Total characters extracted: {char_count:,}")
                        
                        # Create PDF
                        original_filename = uploaded_file.name.replace('.docx', '').replace('.doc', '')
                        pdf_doc = create_pdf_from_text(text_content, original_filename, pdf_options)
                        
                        if pdf_doc:
                            output_filename = f"{original_filename}_converted.pdf"
                            
                            st.download_button(
                                label="📥 Download PDF Document",
                                data=pdf_doc,
                                file_name=output_filename,
                                mime="application/pdf",
                                type="primary",
                                use_container_width=True,
                                key="word_download"
                            )
                            
                            st.success("🎉 Word to PDF conversion completed successfully!")
                            
                            # Show conversion summary
                            with st.expander("📊 Conversion Summary", expanded=True):
                                col1, col2, col3, col4 = st.columns(4)
                                
                                with col1:
                                    st.metric("Paragraphs", len(text_content))
                                
                                with col2:
                                    headings = len([item for item in text_content if item['is_heading']])
                                    st.metric("Headings", headings)
                                
                                with col3:
                                    total_chars = sum(len(item['text']) for item in text_content)
                                    st.metric("Characters", f"{total_chars:,}")
                                
                                with col4:
                                    st.metric("Output Format", "PDF")
                        else:
                            st.error("❌ Failed to create PDF. Please check your Word document and try again.")
                    else:
                        st.warning("⚠️ Could not extract readable content from the Word document.")
                        st.info("💡 **Troubleshooting tips:**")
                        st.write("• Make sure you're uploading a .docx file (not .doc)")
                        st.write("• Ensure the document contains text content")
                        st.write("• Try saving the document in a newer Word format")
                        
                except Exception as e:
                    st.error(f"❌ An error occurred during conversion: {str(e)}")
                    st.info("💡 **Troubleshooting:**")
                    st.write("• Make sure the file is a valid Word document")
                    st.write("• Try using a different Word file")
                    st.write("• Ensure the file isn't corrupted or password-protected")

# Instructions section
if not uploaded_file:
    st.info("👆 **Upload a file to get started!**")
    
    # Mode-specific instructions
    if conversion_mode == "📄➡️📝 PDF to Word":
        st.subheader("📋 PDF to Word Conversion - How it works:")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            **✅ What this conversion does:**
            • Extracts text from PDF files
            • Converts to editable Word documents
            • Preserves basic formatting
            • Adds page numbers and breaks
            • Customizable font settings
            """)
        
        with col2:
            st.markdown("""
            **⚠️ Limitations:**
            • Works best with text-based PDFs
            • Images are not converted
            • Complex layouts may be simplified
            • Scanned PDFs need OCR (not included)
            • Tables may lose formatting
            """)
    
    else:  # Word to PDF
        st.subheader("📋 Word to PDF Conversion - How it works:")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            **✅ What this conversion does:**
            • Extracts text and headings from Word docs
            • Creates professionally formatted PDFs
            • Preserves heading hierarchy
            • Customizable page layout
            • Maintains paragraph structure
            """)
        
        with col2:
            st.markdown("""
            **⚠️ Limitations:**
            • Only supports .docx files
            • Images are not converted
            • Complex formatting may be simplified
            • Tables are not supported
            • Advanced Word features may be lost
            """)
    
    st.subheader("💡 Tips for best results:")
    if conversion_mode == "📄➡️📝 PDF to Word":
        st.markdown("""
        1. **Text-based PDFs work best** - Documents created from Word, Google Docs, etc.
        2. **Avoid scanned documents** - These need OCR processing
        3. **Simple layouts convert better** - Complex multi-column layouts may be reformatted
        4. **Check the preview** - Review extracted text before downloading
        5. **Try different extraction methods** - If one doesn't work well, try the other
        """)
    else:
        st.markdown("""
        1. **Use .docx format** - Older .doc files are not supported
        2. **Simple formatting works best** - Complex layouts may be simplified
        3. **Use built-in heading styles** - These will be preserved in PDF
        4. **Check the preview** - Review extracted content before downloading
        5. **Adjust formatting options** - Customize the PDF appearance to your needs
        """)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666;">
    <small>
    🔧 **Technical Note:** This app works with text-based documents. 
    For scanned documents, OCR software is required. Images and complex layouts may not be perfectly preserved.
    </small>
</div>
""", unsafe_allow_html=True)
